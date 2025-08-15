import streamlit as st
from openai import OpenAI
import PyPDF2
import docx
import requests
from bs4 import BeautifulSoup
import os
from dotenv import load_dotenv
import time
import io

# Load environment variables
load_dotenv()

# Function to get API key from multiple sources
def get_openai_api_key():
    # Try to get from Streamlit secrets first (for cloud deployment)
    try:
        return st.secrets["OPENAI_API_KEY"]
    except:
        pass
    
    # Fallback to environment variable (for local development)
    return os.getenv('OPENAI_API_KEY')

# Initialize session state
if 'generated_questions' not in st.session_state:
    st.session_state.generated_questions = []

class QuestionGenerator:
    def __init__(self):
        api_key = get_openai_api_key()
        if not api_key:
            st.error("❌ OpenAI API key not found!")
            st.stop()
        self.client = OpenAI(api_key=api_key)
    
    def extract_text_from_pdf(self, pdf_file):
        try:
            # Reset file pointer to beginning
            pdf_file.seek(0)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            st.info(f"📖 PDF has {len(reader.pages)} pages")
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n"
                st.info(f"Page {i+1}: {len(page_text)} characters extracted")
            
            if not text.strip():
                st.warning("⚠️ PDF appears to be image-based or encrypted. No text extracted.")
                return ""
            
            st.success(f"✅ PDF: Extracted {len(text)} characters total")
            return text
            
        except Exception as e:
            st.error(f"❌ Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file):
        try:
            # Reset file pointer
            docx_file.seek(0)
            doc = docx.Document(docx_file)
            text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if not text.strip():
                st.warning("⚠️ DOCX file appears to be empty or contains no readable text.")
                return ""
            
            st.success(f"✅ DOCX: Extracted {len(text)} characters from {paragraph_count} paragraphs")
            return text
            
        except Exception as e:
            st.error(f"❌ Error reading DOCX: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, txt_file):
        try:
            # Reset file pointer
            txt_file.seek(0)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    txt_file.seek(0)
                    text = txt_file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text.strip():
                st.warning("⚠️ TXT file appears to be empty.")
                return ""
            
            st.success(f"✅ TXT: Extracted {len(text)} characters")
            return text
            
        except Exception as e:
            st.error(f"❌ Error reading TXT: {str(e)}")
            return ""
    
    def extract_text_from_url(self, url):
        try:
            st.info(f"🌐 Fetching content from: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, timeout=15, headers=headers)
            response.raise_for_status()  # Raise an error for bad status codes
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
                script.extract()
            
            # Try to find main content areas
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            if main_content:
                text = main_content.get_text()
            else:
                text = soup.get_text()
            
            # Clean up the text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                st.warning(f"⚠️ No readable content found at {url}")
                return ""
            
            st.success(f"✅ URL: Extracted {len(text)} characters")
            return text
            
        except requests.exceptions.Timeout:
            st.error(f"❌ Timeout while fetching {url}")
            return ""
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Error fetching URL {url}: {str(e)}")
            return ""
        except Exception as e:
            st.error(f"❌ Error extracting text from URL {url}: {str(e)}")
            return ""
    
    def clean_question_format(self, questions_text):
        """Clean and format the generated questions"""
        lines = questions_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive asterisks and formatting
            line = line.replace('**', '').replace('*', '')
            
            # Clean up question numbering
            if 'Question number:' in line:
                continue  # Skip this line
            
            if 'Question type:' in line:
                # Extract and format question type
                question_type = line.split('Question type:')[-1].strip()
                cleaned_lines.append(f"[Question Type: {question_type}]")
                continue
            
            if 'Difficulty:' in line:
                continue  # Skip difficulty lines
            
            # Clean up the line and add if not empty
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        
        # Join lines and add proper spacing
        result = '\n'.join(cleaned_lines)
        
        # Add proper spacing between questions
        import re
        result = re.sub(r'\n(Q\d+\.)', r'\n\n\1', result)
        
        return result.strip()

    def generate_questions(self, content, difficulty, num_questions, question_types, reference_style):
        difficulty_prompts = {
            "Easy": "Create basic, straightforward questions that test fundamental understanding and recall of key concepts.",
            "Medium": "Create moderate difficulty questions that require analysis, application of concepts, and some critical thinking.",
            "Hard": "Create challenging questions that require synthesis, evaluation, critical analysis, and deep understanding of complex concepts."
        }

        type_list = ", ".join(question_types)

        style_instruction = f"""
        Please use the reference question paper style shown below for formatting and tone, 
        but do NOT copy questions directly. Create new questions with similar style and difficulty.
        
        Reference Question Paper Style:
        {reference_style[:3000]}
        """ if reference_style else ""

        prompt = f"""
        Based on the following content, generate {num_questions} questions of {difficulty} difficulty level.
        
        Difficulty Guidelines: {difficulty_prompts[difficulty]}
        
        Generate ONLY the following question types: {type_list}.

        {style_instruction}
        
        Content for Questions:
        {content[:8000]}
        
        IMPORTANT FORMATTING RULES:
        - Use clean, professional formatting
        - Number questions as: Q1, Q2, Q3, etc.
        - For MCQs: Use (A), (B), (C), (D) format and clearly mark correct answer
        - For Short/Long Answer: Just state the question clearly
        - For Fill in the Blanks: Use clear _______ blanks
        - For True/False: State clearly and ask for justification
        - Add appropriate spacing between questions
        - No asterisks or special formatting symbols
        
        Example format:
        
        Q1. [Question Type: MCQ]
        What is the primary purpose of delegation in management?
        
        (A) To reduce manager workload only
        (B) To develop employee skills and improve efficiency
        (C) To avoid responsibility
        (D) To create confusion in the organization
        
        Correct Answer: (B)
        
        Q2. [Question Type: Short Answer]
        Define the concept of delegation in management and explain its key benefits.
        
        Q3. [Question Type: Fill in the Blanks]
        Effective delegation requires clear _______ of tasks and proper _______ mechanisms.
        
        Generate exactly {num_questions} questions following this clean format.
        """

        try:
            st.info("🤖 Calling OpenAI API...")
            
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert educator who creates high-quality examination questions."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=3000,
                temperature=0.7
            )
            
            result = response.choices[0].message.content
            
            # Clean and format the questions
            cleaned_result = self.clean_question_format(result)
            
            st.success("✅ Questions generated and formatted successfully!")
            return cleaned_result
            
        except Exception as e:
            st.error(f"❌ Error generating questions: {str(e)}")
            return ""

def main():
    st.set_page_config(
        page_title="AI Question Bank Generator",
        page_icon="📝",
        layout="wide"
    )
    
    st.title("🎓 AI Question Bank Generator")
    st.markdown("Upload documents and generate customized question banks based on your selection.")
    
    # Check API key
    api_key = get_openai_api_key()
    if not api_key:
        st.error("❌ OpenAI API key not found!")
        st.markdown("""
        **For local development:** Add it to your .env file:
        ```
        OPENAI_API_KEY=your_key_here
        ```
        
        **For Streamlit Cloud:** Add it to your app's secrets in the dashboard.
        
        You can get your API key from: https://platform.openai.com/account/api-keys
        """)
        st.stop()
    
    qg = QuestionGenerator()
    
    with st.sidebar:
        st.header("⚙️ Settings")
        st.success("✅ API Key Loaded")
        
        # Debug information
        st.subheader("🐛 Debug Info")
        if st.button("Show API Key Status"):
            if api_key:
                st.success(f"✅ API Key loaded (starts with: {api_key[:10]}...)")
            else:
                st.error("❌ No API Key found")
        
        if st.button("Show Environment"):
            st.write("Environment variables loaded:", bool(load_dotenv()))
            st.write("API Key present:", bool(api_key))
        
        if st.button("🗑️ Clear Questions"):
            st.session_state.generated_questions = []
            st.rerun()
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📄 Upload Study Materials")
        uploaded_files = st.file_uploader(
            "Choose PDF, DOCX, or TXT files (max 2)",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="Upload your study materials"
        )
        
        if uploaded_files and len(uploaded_files) > 2:
            st.warning("Maximum 2 files allowed")
            uploaded_files = uploaded_files[:2]
        
        st.header("📄 (Optional) Reference Question Paper")
        ref_file = st.file_uploader(
            "Upload reference paper (PDF, DOCX, or TXT)",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=False,
            help="Upload a sample question paper to match the style"
        )
        
        st.header("🔗 Web Links (Optional)")
        url1 = st.text_input("URL 1", placeholder="https://example.com/article")
        url2 = st.text_input("URL 2", placeholder="https://example.com/resource")
        
        st.header("⚙️ Question Settings")
        col_set1, col_set2 = st.columns(2)
        
        with col_set1:
            difficulty = st.selectbox("Difficulty", ["Easy", "Medium", "Hard"], index=1)
        
        with col_set2:
            num_questions = st.slider("Questions", 5, 50, 15, 5)
        
        question_types = st.multiselect(
            "Select Question Types",
            ["Multiple Choice Questions (MCQ)", "Short Answer", "Long Answer", "Fill in the Blanks", "True/False"],
            default=["Multiple Choice Questions (MCQ)", "Short Answer"],
            help="Choose the types of questions you want to generate"
        )
        
        # Generate button
        generate_disabled = not uploaded_files and not url1.strip() and not url2.strip()
        
        if st.button("🚀 Generate Questions", type="primary", use_container_width=True, disabled=generate_disabled):
            if generate_disabled:
                st.error("❌ Please upload at least one file or provide a URL")
            else:
                # Show processing status
                with st.container():
                    st.subheader("📋 Processing Status")
                    
                    # Progress tracking
                    progress_placeholder = st.empty()
                    status_placeholder = st.empty()
                    
                    all_content = ""
                    total_sources = len([f for f in uploaded_files if f] if uploaded_files else [])
                    total_sources += len([u for u in [url1, url2] if u.strip()])
                    
                    processed = 0
                    
                    # Process uploaded files
                    if uploaded_files:
                        for file in uploaded_files:
                            status_placeholder.info(f"Processing {file.name}...")
                            
                            if file.type == "application/pdf":
                                text = qg.extract_text_from_pdf(file)
                            elif "wordprocessingml" in file.type:
                                text = qg.extract_text_from_docx(file)
                            elif file.type == "text/plain":
                                text = qg.extract_text_from_txt(file)
                            else:
                                st.warning(f"⚠️ Unsupported file type: {file.type}")
                                text = ""
                            
                            if text.strip():
                                all_content += f"\n\n=== Content from {file.name} ===\n{text}"
                            
                            processed += 1
                            progress_placeholder.progress(processed / total_sources)
                    
                    # Process URLs
                    for i, url in enumerate([url1, url2], 1):
                        if url.strip():
                            status_placeholder.info(f"Processing URL {i}...")
                            text = qg.extract_text_from_url(url.strip())
                            
                            if text.strip():
                                all_content += f"\n\n=== Content from URL {i} ===\n{text}"
                            
                            processed += 1
                            progress_placeholder.progress(processed / total_sources)
                    
                    # Process reference file
                    reference_style = ""
                    if ref_file:
                        status_placeholder.info("Processing reference paper...")
                        
                        if ref_file.type == "application/pdf":
                            reference_style = qg.extract_text_from_pdf(ref_file)
                        elif "wordprocessingml" in ref_file.type:
                            reference_style = qg.extract_text_from_docx(ref_file)
                        elif ref_file.type == "text/plain":
                            reference_style = qg.extract_text_from_txt(ref_file)
                        
                        if reference_style.strip():
                            st.info("📄 Reference paper style will be used for question formatting.")
                    
                    # Check if we have content
                    if all_content.strip():
                        st.success(f"✅ Successfully extracted {len(all_content)} characters from all sources")
                        
                        # Generate questions
                        status_placeholder.info("Generating questions with AI...")
                        questions = qg.generate_questions(all_content, difficulty, num_questions, question_types, reference_style)
                        
                        if questions and questions.strip():
                            st.session_state.generated_questions = questions
                            st.success(f"🎉 Generated {num_questions} questions successfully!")
                        else:
                            st.error("❌ Failed to generate questions. Please try again.")
                    else:
                        st.error("❌ No content could be extracted from any source. Please check your files and URLs.")
                        
                        # Show debugging info
                        st.subheader("🐛 Debug Information")
                        st.write(f"Files uploaded: {len(uploaded_files) if uploaded_files else 0}")
                        st.write(f"URLs provided: {len([u for u in [url1, url2] if u.strip()])}")
                        
                        if uploaded_files:
                            for file in uploaded_files:
                                st.write(f"- {file.name}: {file.size} bytes, type: {file.type}")
                    
                    # Clear progress indicators
                    progress_placeholder.empty()
                    status_placeholder.empty()
    
    with col2:
        st.header("📊 Summary")
        
        if uploaded_files:
            st.subheader("📁 Files")
            for file in uploaded_files:
                file_size = file.size / 1024  # Convert to KB
                st.write(f"• {file.name} ({file_size:.1f} KB)")
        
        if ref_file:
            st.subheader("📄 Reference Paper")
            st.write(f"• {ref_file.name}")
        
        urls = [url for url in [url1, url2] if url.strip()]
        if urls:
            st.subheader("🌐 URLs")
            for i, url in enumerate(urls, 1):
                st.write(f"• URL {i}: {url[:30]}{'...' if len(url) > 30 else ''}")
        
        st.subheader("⚙️ Settings")
        if 'difficulty' in locals():
            st.write(f"**Difficulty:** {difficulty}")
            st.write(f"**Questions:** {num_questions}")
            if 'question_types' in locals() and question_types:
                st.write(f"**Types:** {', '.join(question_types[:2])}{'...' if len(question_types) > 2 else ''}")
    
    # Display results
    if st.session_state.generated_questions:
        st.header("📝 Generated Question Bank")
        
        col_d1, col_d2 = st.columns(2)
        
        with col_d1:
            st.download_button(
                "📄 Download TXT",
                data=st.session_state.generated_questions,
                file_name=f"question_bank_{difficulty}_{num_questions}_{int(time.time())}.txt",
                mime="text/plain"
            )
        
        with col_d2:
            formatted = f"""
QUESTION BANK
==============
Difficulty: {difficulty if 'difficulty' in locals() else 'Medium'}
Questions: {num_questions if 'num_questions' in locals() else 25}
Types: {', '.join(question_types) if 'question_types' in locals() and question_types else 'MCQ, Short Answer'}
Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}

{st.session_state.generated_questions}
"""
            st.download_button(
                "📋 Download Formatted",
                data=formatted,
                file_name=f"question_bank_formatted_{int(time.time())}.txt",
                mime="text/plain"
            )
        
        # Show character count
        char_count = len(st.session_state.generated_questions)
        st.caption(f"Generated {char_count} characters")
        
        st.text_area(
            "Questions:",
            value=st.session_state.generated_questions,
            height=400
        )

if __name__ == "__main__":
    main()